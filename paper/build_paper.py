"""Build the NyxChat paper from paper/nyxchat.md into DOCX and LaTeX.

Usage:  python paper/build_paper.py
Produces: paper/NyxChat_Paper.docx and paper/nyxchat.tex

Markdown subset used by the paper:
  # Title (first line), > authors line, ## Section, ### Subsection,
  paragraphs, "- " bullets, pipe tables (+ optional "Table: caption" line),
  **bold**, *italic*, `code`, [n] citations, ![caption](path) figures,
  a final "## References" section with "[n] text" paragraphs.
"""
import re
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parent
SRC = ROOT.joinpath("nyxchat.md")


def parse(md):
    lines = md.splitlines()
    title = lines[0].lstrip("# ").strip()
    blocks = []
    idx = 1
    para = []

    def flush():
        nonlocal para
        if para:
            blocks.append(("p", " ".join(s.strip() for s in para)))
            para = []

    while idx < len(lines):
        line = lines[idx]
        if line.startswith("> "):
            flush(); blocks.append(("authors", line[2:].strip()))
        elif line.startswith("### "):
            flush(); blocks.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            flush(); blocks.append(("h2", line[3:].strip()))
        elif line.startswith("- "):
            flush()
            items = []
            while idx < len(lines) and lines[idx].startswith("- "):
                items.append(lines[idx][2:].strip()); idx += 1
            blocks.append(("ul", items)); continue
        elif line.startswith("|"):
            flush()
            rows = []
            while idx < len(lines) and lines[idx].startswith("|"):
                cells = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                idx += 1
            caption = None
            if idx < len(lines) and lines[idx].startswith("Table:"):
                caption = lines[idx][6:].strip(); idx += 1
            blocks.append(("table", rows, caption)); continue
        elif line.startswith("!["):
            flush()
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            blocks.append(("fig", m.group(1), m.group(2)))
        elif line.strip() == "":
            flush()
        else:
            para.append(line)
        idx += 1
    flush()
    return title, blocks


def add_runs(paragraph, text):
    pos = 0
    for m in re.finditer(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)", text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(2):
            paragraph.add_run(m.group(2)).bold = True
        elif m.group(3):
            paragraph.add_run(m.group(3)).italic = True
        else:
            run = paragraph.add_run(m.group(4)); run.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(title, blocks, out):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_no = 0
    fig_no = 0
    for b in blocks:
        kind = b[0]
        if kind == "authors":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_runs(p, b[1])
        elif kind == "h2":
            doc.add_heading(b[1], level=1)
        elif kind == "h3":
            doc.add_heading(b[1], level=2)
        elif kind == "p":
            p = doc.add_paragraph(); add_runs(p, b[1]); p.paragraph_format.space_after = Pt(6)
        elif kind == "ul":
            for item in b[1]:
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, item)
        elif kind == "table":
            rows, caption = b[1], b[2]
            table_no += 1
            if caption:
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cp.add_run("Table %d. %s" % (table_no, caption)); run.italic = True; run.font.size = Pt(9)
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for rowIndex, row in enumerate(rows):
                for colIndex, cell in enumerate(row):
                    c = t.cell(rowIndex, colIndex); c.text = ""
                    add_runs(c.paragraphs[0], cell)
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        if rowIndex == 0:
                            run.bold = True
            doc.add_paragraph()
        elif kind == "fig":
            fig_no += 1
            path = ROOT.joinpath(b[2])
            if path.exists():
                doc.add_picture(str(path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cp.add_run("Figure %d. %s" % (fig_no, b[1])); run.italic = True; run.font.size = Pt(9)
    doc.save(out)


def tex_escape(s):
    s = s.replace("\\", "\\textbackslash{}")
    for ch in "&%$#_{}":
        s = s.replace(ch, "\\" + ch)
    return s.replace("~", "\\textasciitilde{}").replace("^", "\\textasciicircum{}")


def tex_inline(text):
    out = ""
    pos = 0
    for m in re.finditer(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(\d+(?:,\s*\d+)*)\])", text):
        out += tex_escape(text[pos:m.start()])
        if m.group(2):
            out += "\\textbf{" + tex_escape(m.group(2)) + "}"
        elif m.group(3):
            out += "\\emph{" + tex_escape(m.group(3)) + "}"
        elif m.group(4):
            out += "\\texttt{" + tex_escape(m.group(4)) + "}"
        else:
            out += "\\cite{" + ",".join("ref" + n.strip() for n in m.group(5).split(",")) + "}"
        pos = m.end()
    return out + tex_escape(text[pos:])


def build_tex(title, blocks, out):
    head = ["\\documentclass[conference]{IEEEtran}",
            "\\usepackage{booktabs,graphicx,url,amsmath}",
            "\\begin{document}",
            "\\title{" + tex_escape(title) + "}"]
    body, refs = [], []
    in_refs = False
    abstract = None
    for b in blocks:
        kind = b[0]
        if kind == "authors":
            head.append("\\author{" + tex_inline(b[1]).replace(";", " \\and ") + "}")
            head.append("\\maketitle"); continue
        if kind == "h2" and b[1].lower() == "references":
            in_refs = True; continue
        if in_refs:
            if kind == "p":
                for m in re.finditer(r"\[(\d+)\]\s*(.+?)(?=\s\[\d+\]|$)", b[1]):
                    refs.append((m.group(1), m.group(2)))
            continue
        if kind == "h2" and b[1].lower() == "abstract":
            abstract = "pending"; continue
        if abstract == "pending" and kind == "p":
            head.append("\\begin{abstract}" + tex_inline(b[1]) + "\\end{abstract}")
            abstract = "done"; continue
        if kind == "h2":
            body.append("\\section{" + tex_escape(b[1]) + "}")
        elif kind == "h3":
            body.append("\\subsection{" + tex_escape(b[1]) + "}")
        elif kind == "p":
            body.append(tex_inline(b[1]) + "\n")
        elif kind == "ul":
            body.append("\\begin{itemize}")
            body += ["  \\item " + tex_inline(i) for i in b[1]]
            body.append("\\end{itemize}")
        elif kind == "table":
            rows, caption = b[1], b[2]
            body.append("\\begin{table}[t]\\centering\\footnotesize")
            if caption:
                body.append("\\caption{" + tex_inline(caption) + "}")
            body.append("\\begin{tabular}{" + "l" * len(rows[0]) + "}\\toprule")
            body.append(" & ".join(tex_inline(c) for c in rows[0]) + " \\\\ \\midrule")
            for r in rows[1:]:
                body.append(" & ".join(tex_inline(c) for c in r) + " \\\\")
            body.append("\\bottomrule\\end{tabular}\\end{table}")
        elif kind == "fig":
            body.append("\\begin{figure}[t]\\centering")
            body.append("\\includegraphics[width=\\columnwidth]{" + b[2] + "}")
            body.append("\\caption{" + tex_inline(b[1]) + "}\\end{figure}")
    parts = head + body
    if refs:
        parts.append("\\begin{thebibliography}{" + str(len(refs)) + "}")
        for n, t in refs:
            parts.append("\\bibitem{ref" + n + "} " + tex_inline(t))
        parts.append("\\end{thebibliography}")
    parts.append("\\end{document}")
    Path(out).write_text("\n".join(parts) + "\n", encoding="utf-8")


if __name__ == "__main__":
    md = SRC.read_text(encoding="utf-8")
    title, blocks = parse(md)
    build_docx(title, blocks, ROOT.joinpath("NyxChat_Paper.docx"))
    build_tex(title, blocks, ROOT.joinpath("nyxchat.tex"))
    print("built NyxChat_Paper.docx and nyxchat.tex in", ROOT)